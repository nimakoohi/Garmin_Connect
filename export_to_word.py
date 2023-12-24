import os
from docx import Document
from docx.shared import Pt, RGBColor


def add_formatted_paragraph(doc, text, bold=True, color=None, font_size=None):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_size:
        run.font.size = Pt(font_size)
    return paragraph


def running_evaluation_questions(times, types_of_runs,
                                 res_Speed, res_hr, res_cd, res_al,
                                 standard_deviation_pace,
                                 standard_deviation_heart_rate,
                                 standard_deviation_cadence,
                                 standard_deviation_altitude, save_directory,
                                 name, distances):
    Physical_Comfort_Response1 = input('''How did your body feel during the run?
                                       Any specific areas of discomfort or pain? ''')
    Physical_Comfort_Response2 = input('''Did you experience any muscle tightness, soreness, or stiffness? ''')

    Breathing_and_Cardiovascular_Response1 = input('''How was your breathing throughout the run? Was it comfortable or labored? ''')
    Breathing_and_Cardiovascular_Response2 = input('''Did you notice any
                                                   changes in your heart rate, and if so, were they within a comfortable range? ''')

    Energy_Levels_Response = input('''How would you describe your energy
                                   levels during the run? Did you feel fatigued at any point? ''')

    Mental_State_Response1 = input('''What was your mental state like during
                                   the run? Were you focused, distracted, or perhaps in a flow state? ''')

    Mental_State_Response2 = input('''Did you experience any mental barriers or breakthroughs? ''')

    Running_Form_Response1 = input('''Were you mindful of your running form? Did you notice any changes in your gait or posture? ''')
    Running_Form_Response2 = input('''Did you encounter any challenges related to your form? ''')

    Terrain_and_Environment_Response1 = input('''How did the terrain (e.g., flat, hilly, uneven) impact your running experience? ''')
    Terrain_and_Environment_Response2 = input('''Did the weather or environmental conditions affect your performance or enjoyment? ''')

    Hydration_and_Nutrition_Response1 = input('''Did you feel adequately hydrated and fueled before and during the run? ''')
    Hydration_and_Nutrition_Response2 = input('''Did you experience any issues related to nutrition or hydration? ''')

    Goal_Achievement_Response1 = input('''Were you able to meet the goals you set for yourself during the run? ''')
    Goal_Achievement_Response2 = input('''How do you feel about your overall performance and progress? ''')

    Recovery_Response1 = input('''How is your body feeling post-run? Any lingering discomfort or signs of quick recovery? ''')
    Recovery_Response2 = input('''Did you engage in any post-run stretching or recovery activities? ''')

    Overall_Satisfaction1 = input('''On a scale of 1 to 10, how satisfied are you with your running experience today? ''')
    Overall_Satisfaction2 = input('''What aspects of the run brought you the most joy or fulfillment? ''')

    # Create a new Word document
    doc = Document()

    # Add content to the document 
    doc.add_heading(f'{(distances[-1])[:2]} kilometers on {times[0]}', level=1)

    # Add content to the document
    add_formatted_paragraph(doc, "Hi.\nI hope this message finds you well. I wanted to share the details of my recent "
                                 f"{types_of_runs} for your analysis and feedback.\n\n**Run Details:**\n- Distance: "
                                 f"{(distances[-1])[:2]} kilometers\n- Duration: {times[0]} to {times[-1]}\n\n**Average Pace "
                                 f"per Kilometer:**\n- {res_Speed}\n\n**Average Heart Rate per Kilometer:**\n- {res_hr}\n\n"
                                 f"**Average Cadence per Kilometer:**\n- {res_cd}\n\n**Average Altitude per Kilometer:**\n- "
                                 f"{res_al}\n\n**Standard Deviation:**\n- Pace: {standard_deviation_pace} m/s\n- Heart Rate: "
                                 f"{standard_deviation_heart_rate} bpm\n- Cadence: {standard_deviation_cadence} steps/min\n- "
                                 f"Altitude: {standard_deviation_altitude} m\n\nI would greatly appreciate your insights and "
                                 "any recommendations you may have based on this data. Thank you in advance for your guidance."
                                 "\n\nBest regards,\n\nNima\n", bold=False, color=None, font_size=14)

    add_formatted_paragraph(doc, "These are my answers after the run:", bold=True, color=(0, 0, 0), font_size=18)  # Black color

    # Physical Comfort section
    add_formatted_paragraph(doc, "Physical Comfort:", bold=True, color=(0, 0, 255), font_size=16)  # Blue color
    add_formatted_paragraph(doc, "1. How did your body feel during the run? Any specific areas of discomfort or pain?",
                            font_size=14)
    add_formatted_paragraph(doc, Physical_Comfort_Response1)
    add_formatted_paragraph(doc, "2. Did you experience any muscle tightness, soreness, or stiffness?", font_size=14)
    add_formatted_paragraph(doc, Physical_Comfort_Response2)

   
    # Breathing and Cardiovascular Response:
    add_formatted_paragraph(doc, "Breathing and Cardiovascular Response:", bold=True, color=(0, 0, 255), font_size=16)  # Blue color
    add_formatted_paragraph(doc, "1. How was your breathing throughout the run? Was it comfortable or labored?", font_size=14)
    add_formatted_paragraph(doc, Breathing_and_Cardiovascular_Response1)
    add_formatted_paragraph(doc, "2. Did you notice any changes in your heart rate, and if so, were they within a comfortable range?",
                            font_size=14)
    add_formatted_paragraph(doc, (Breathing_and_Cardiovascular_Response2))

    #Energy Levels:
    add_formatted_paragraph(doc, "Energy Levels:", bold=True, color=(0, 0, 255), font_size=16)  # Blue color
    add_formatted_paragraph(doc, "1. How would you describe your energy levels during the run? Did you feel fatigued at any point?",
                            font_size=14)
    add_formatted_paragraph(doc, (Energy_Levels_Response))

    #Mental State:
    add_formatted_paragraph(doc, "Mental State:", bold=True, color=(0, 0, 255), font_size=16)  # Blue color
    add_formatted_paragraph(doc, "1. What was your mental state like during the run? Were you focused, distracted, or perhaps in a flow state?",
                            font_size=14)
    add_formatted_paragraph(doc, (Mental_State_Response1))
    add_formatted_paragraph(doc, "2. Did you experience any mental barriers or breakthroughs?", font_size=14)
    add_formatted_paragraph(doc, (Mental_State_Response2))

    #Running Form:
    add_formatted_paragraph(doc, "Running Form:", bold=True, color=(0, 0, 255), font_size=16)  # Blue color
    add_formatted_paragraph(doc, "1. Were you mindful of your running form? Did you notice any changes in your gait or posture?",
                            font_size=14)
    add_formatted_paragraph(doc, (Running_Form_Response1))
    add_formatted_paragraph(doc, "2. Did you encounter any challenges related to your form?",
                            font_size=14)
    add_formatted_paragraph(doc, (Running_Form_Response2))

    #Terrain and Environment:
    add_formatted_paragraph(doc, "Terrain and Environment:", bold=True, color=(0, 0, 255), font_size=16)  # Blue color
    add_formatted_paragraph(doc, "1. How did the terrain (e.g., flat, hilly, uneven) impact your running experience?", font_size=14)
    add_formatted_paragraph(doc, Terrain_and_Environment_Response1)
    add_formatted_paragraph(doc, "2. Did the weather or environmental conditions affect your performance or enjoyment?", font_size=14)
    add_formatted_paragraph(doc, Terrain_and_Environment_Response2)

    #Hydration and Nutrition:
    add_formatted_paragraph(doc, "Hydration and Nutrition:", bold=True, color=(0, 0, 255), font_size=16)  # Blue color
    add_formatted_paragraph(doc, "1. Did you feel adequately hydrated and fueled before and during the run?", font_size=14)
    add_formatted_paragraph(doc, Hydration_and_Nutrition_Response1)
    add_formatted_paragraph(doc, "2. Did you experience any issues related to nutrition or hydration?", font_size=14)
    add_formatted_paragraph(doc, Hydration_and_Nutrition_Response2)

    #Goal Achievement:
    add_formatted_paragraph(doc, "Goal Achievement:", bold=True, color=(0, 0, 255), font_size=16)  # Blue color
    add_formatted_paragraph(doc, "1. Were you able to meet the goals you set for yourself during the run?", font_size=14)
    add_formatted_paragraph(doc, Goal_Achievement_Response1)
    add_formatted_paragraph(doc, "2. How do you feel about your overall performance and progress?", font_size=14)
    add_formatted_paragraph(doc, Goal_Achievement_Response2)

    #Recovery:
    add_formatted_paragraph(doc, "Recovery:", bold=True, color=(0, 0, 255), font_size=16)  # Blue color
    add_formatted_paragraph(doc, "1. How is your body feeling post-run? Any lingering discomfort or signs of quick recovery?", font_size=14)
    add_formatted_paragraph(doc, Recovery_Response1)
    add_formatted_paragraph(doc, "2. Did you engage in any post-run stretching or recovery activities?", font_size=14)
    add_formatted_paragraph(doc, Recovery_Response2)

    #Overall Satisfaction:
    add_formatted_paragraph(doc, "Overall Satisfaction:", bold=True, color=(0, 0, 255), font_size=16)  # Blue color
    add_formatted_paragraph(doc, "1. On a scale of 1 to 10, how satisfied are you with your running experience today?", font_size=14)
    add_formatted_paragraph(doc, Overall_Satisfaction1)
    add_formatted_paragraph(doc, "2. What aspects of the run brought you the most joy or fulfillment?", font_size=14)
    add_formatted_paragraph(doc, Overall_Satisfaction2)

    # Save the document
    doc_file_path = os.path.join(save_directory, f'{name} - {types_of_runs} - {times[0].split()[0]} - {(distances[-1])[:2]}km.docx')
    doc.save(doc_file_path)

    print("Output has been printed to Word document:", doc_file_path)