# 1. Add all wourkout on Database
# 2. Compare long run with each other
# 3. Link to the site
# 4. Add other plot
# 5. get a tcx files in folder and extract word and pictures into another folder.

import os
import pandas as pd
from docx import Document
from DatabaseManager import DatabaseManager
from TCXParser import TCXParser
from Plotting import Plotting
from ConfigHandler import read_config

def connect_to_database(file_path, name, types_of_runs):
    config = read_config()

    db = DatabaseManager(config)

    (times, distances, heart_rates, speeds,
     cadances, altitude_meters) = TCXParser.parse(file_path)

    time_column = times
    heart_rate_column = heart_rates
    speed_column = speeds
    cadence_column = cadances
    altitude_column = altitude_meters

    if db.establish_connection():
        insert_query_run_name = ("""
                                 INSERT INTO Run_Name(run_name) VALUES (%s)
                                 """)

        run_name = f'{name}_{types_of_runs}_{times[0].split()[0]}'

        db.cursor.execute(insert_query_run_name, (run_name, ))

        db.cursor.execute(f"SELECT * FROM Run_Name WHERE run_name = '{run_name}'")
        run_name_id = db.cursor.fetchone()[0]
        run_name_id = [run_name_id]

        run_name_id.extend([run_name_id[0]] * (len(times) - 1))
        
        data = list(zip(run_name_id, times, distances, heart_rates, speeds,
                    cadances, altitude_meters))

        insert_query = ("""
        INSERT INTO ActivityData (run_name_id, Time, Distance, Heart_Rate,
        Speed, Cadance, Altitude_Meters) VALUES (%s, %s, %s, %s, %s, %s, %s)
        """)

        db.cursor.executemany(insert_query, data)
        db.connection.commit()
    else:
        print("Failed to establish connection")
        
    return (times, time_column, heart_rate_column, speed_column,
            cadence_column, altitude_column, db)

def create_directory(directory_path):
    try:
        # Create target directory
        os.makedirs(directory_path)
        print(f"Directory '{directory_path}' created successfully.")
    except FileExistsError:
        print(f"Directory '{directory_path}' already exists.")


def main():
    """
        The main function of the program. It initializes the database connection,
        retrieves the TCX files from the specified directory, and stores the data
        into the database. The total distance traveled and the average speed for each
        km of the day are calculated and printed to the console.
    """

    res_Speed = []
    res_hr = []
    res_cd = []
    res_al = []

    save_directory = os.path.join(os.getcwd(),
                                  input("Please insert the directory where you want store your files: "))
    create_directory(save_directory)
    
    file_path = input('Please insert the directory where tcx files are stored: ')
    while not os.path.exists(file_path):
        print("Error: The specified directory does not exist.")
        file_path = input('Please insert the directory where tcx files are stored: ')

    name = input("Please insert your full name: ")

    types_of_runs = input(''' Please tell us what types of runs you do:
    - Recovery Run
    - Long Run
    - Interval Run
    - Tempo Run
    - Fartlek Run
    - Hill Run
    - Progression Run
    - Cross-Training Run 
    ''')

    (times, time_column, heart_rate_column, speed_column,
     cadence_column, altitude_column, db) = connect_to_database(file_path, name, types_of_runs)

    find_total_km = "SELECT Distance AS row_count FROM ActivityData ORDER BY id DESC LIMIT 1;"

    db.cursor.execute(find_total_km)
    total_km_result = db.cursor.fetchone()[0]

    if total_km_result:
        total_km = round(total_km_result/1000)
        total_km_extra = total_km_result - (total_km*1000)
    else:
        print("Failed to retrieve Total km.")

    Plotting(time_column, heart_rate_column, speed_column, cadence_column,
             altitude_column, save_directory, name, total_km_result, types_of_runs, times)

    find_total_km = '''SELECT Distance AS row_count FROM ActivityData
    ORDER BY id DESC LIMIT 1'''

    db.cursor.execute(find_total_km)
    total_km_result = round(db.cursor.fetchone()[0])

    # Create a pandas dataframe
    df = pd.DataFrame({
        "Pace (m/s)": speed_column,
        "Heart Rate (bpm)": heart_rate_column,
        "Cadence (steps/min)": cadence_column,
        "Altitude (m)": altitude_column
    })

    # Convert 'Pace (m/s)' column to numeric values
    df['Pace (m/s)'] = pd.to_numeric(df['Pace (m/s)'], errors='coerce')

    # Calculate the standard deviation of pace
    standard_deviation_pace = df['Pace (m/s)'].std()

    df['Heart Rate (bpm)'] = pd.to_numeric(df['Heart Rate (bpm)'], errors='coerce')
    # Calculate the standard deviation of heart rate
    standard_deviation_heart_rate = df['Heart Rate (bpm)'].std()

    df['Cadence (steps/min)'] = pd.to_numeric(df['Cadence (steps/min)'], errors='coerce')
    # Calculate the standard deviation of cadence
    standard_deviation_cadence = df['Cadence (steps/min)'].std()

    df['Altitude (m)'] = pd.to_numeric(df['Altitude (m)'], errors='coerce')
    # Calculate the standard deviation of altitude
    standard_deviation_altitude = df['Altitude (m)'].std()

    try:
        for find_another_km in range(0, total_km):
            another_km = f'''SELECT
            CONCAT(
                CASE
                    WHEN LENGTH(FLOOR(60 / (AVG(Speed) * 3.6))) = 1 THEN '0'
                    ELSE ''
                END,
                FLOOR(60 / (AVG(Speed) * 3.6)),
                ':',
                LPAD(CEILING(((60 / (AVG(Speed) * 3.6)) - FLOOR(60 / (AVG(Speed) * 3.6))) * 60), 2, '0')
            ) AS result
            FROM ActivityData
            WHERE Distance BETWEEN {find_another_km}000
            AND {find_another_km+1}000
            '''
            
            db.connection.commit()
            db.cursor.execute(another_km)
            another_km = db.cursor.fetchone()[0]
            res_Speed.append(another_km)
            
            hr = f'''
            SELECT CEILING(AVG(Heart_Rate)) AS HR100 FROM activitydata
            WHERE Distance BETWEEN {find_another_km}000 AND {find_another_km+1}000 ORDER BY id DESC LIMIT 1;
            '''
            
            db.connection.commit()
            db.cursor.execute(hr)
            hr = db.cursor.fetchone()[0]
            res_hr.append(hr)
            
            cd = f'''
            SELECT CEILING(AVG(Cadance)) AS cd FROM activitydata
            WHERE Distance BETWEEN {find_another_km}000 AND {find_another_km+1}000 ORDER BY id DESC LIMIT 1;
            '''
            
            db.connection.commit()
            db.cursor.execute(cd)
            cd = db.cursor.fetchone()[0]
            res_cd.append(cd)

            al_me = f'''
            SELECT CEILING(AVG(Altitude_Meters)) AS AL_Me FROM activitydata
            WHERE Distance BETWEEN {find_another_km}000 AND {find_another_km+1}000 ORDER BY id DESC LIMIT 1;
            '''

            db.connection.commit()
            db.cursor.execute(al_me)
            al_me = db.cursor.fetchone()[0]
            res_al.append(al_me)

        if total_km_extra > 0:
            last_km = f'''SELECT
            CONCAT(
                CASE
                    WHEN LENGTH(FLOOR(60 / (AVG(Speed) * 3.6))) = 1 THEN '0'
                    ELSE ''
                END,
                FLOOR(60 / (AVG(Speed) * 3.6)),
                ':',
                LPAD(CEILING(((60 / (AVG(Speed) * 3.6)) - FLOOR(60 / (AVG(Speed) * 3.6))) * 60), 2, '0')
            ) AS result
            FROM ActivityData
            WHERE Distance BETWEEN {total_km*1000} AND {total_km_result} '''

            db.connection.commit()
            db.cursor.execute(last_km)
            last_km = db.cursor.fetchone()[0]

            res_Speed.append(last_km)

            hr = f'''
            SELECT CEILING(AVG(Heart_Rate)) AS HR100 FROM activitydata
            WHERE Distance BETWEEN {total_km*1000} AND {total_km_result} ORDER BY id DESC LIMIT 1;
            '''

            db.connection.commit()
            db.cursor.execute(hr)
            hr = db.cursor.fetchone()[0]
            res_hr.append(hr)

            cd = f'''
            SELECT CEILING(AVG(Cadance)) AS cd FROM activitydata
            WHERE Distance BETWEEN {total_km*1000} AND {total_km_result} ORDER BY id DESC LIMIT 1;
            '''

            db.connection.commit()
            db.cursor.execute(cd)
            cd = db.cursor.fetchone()[0]
            res_cd.append(cd)

            al_me = f'''
            SELECT CEILING(AVG(Altitude_Meters)) AS AL_Me FROM activitydata
            WHERE Distance BETWEEN {total_km*1000} AND {total_km_result} ORDER BY id DESC LIMIT 1;
            '''

            db.connection.commit()
            db.cursor.execute(al_me)
            al_me = db.cursor.fetchone()[0]
            res_al.append(al_me)

    except mysql.connector.Error as e:
        print(f"Database error: {e.errno} - {e.msg}")
    except Exception as e:
        print(f"Other error: {e}")

    db.close_connection()

    Physical_Comfort_Response1 = input('''How did your body feel during the run?
                                       Any specific areas of discomfort or pain? ''')
    Physical_Comfort_Response2 = input('''Did you experience any muscle tightness, soreness, or stiffness? ''')

    Breathing_and_Cardiovascular_Response1 = input('''How was your breathing throughout the run? Was it comfortable or labored? ''')
    Breathing_and_Cardiovascular_Response2 = input('''Did you notice any
                                                   changes in your heart rate, and if so, were they within a comfortable range? ''')

    Energy_Levels_Response = input('''How would you describe your energy
                                   levels during the run? Did you feel fatigued at any point? ''')

    Mental_State_Response1 = input('''What was your mental state like during
                                   the run? Were you focused, distracted, or perhaps in a flow state? ''')

    Mental_State_Response2 = input('''Did you experience any mental barriers or breakthroughs? ''')

    Running_Form_Response1 = input('''Were you mindful of your running form? Did you notice any changes in your gait or posture? ''')
    Running_Form_Response2 = input('''Did you encounter any challenges related to your form? ''')

    Terrain_and_Environment_Response1 = input('''How did the terrain (e.g., flat, hilly, uneven) impact your running experience? ''')
    Terrain_and_Environment_Response2 = input('''Did the weather or environmental conditions affect your performance or enjoyment? ''')

    Hydration_and_Nutrition_Response1 = input('''Did you feel adequately hydrated and fueled before and during the run? ''')
    Hydration_and_Nutrition_Response2 = input('''Did you experience any issues related to nutrition or hydration? ''')

    Goal_Achievement_Response1 = input('''Were you able to meet the goals you set for yourself during the run? ''')
    Goal_Achievement_Response2 = input('''How do you feel about your overall performance and progress? ''')

    Recovery_Response1 = input('''How is your body feeling post-run? Any lingering discomfort or signs of quick recovery? ''')
    Recovery_Response2 = input('''Did you engage in any post-run stretching or recovery activities? ''')

    Overall_Satisfaction1 = input('''On a scale of 1 to 10, how satisfied are you with your running experience today? ''')
    Overall_Satisfaction2 = input('''What aspects of the run brought you the most joy or fulfillment? ''')

    # Create a new Word document
    doc = Document()

    # Add content to the document 
    doc.add_heading(f'{total_km_result} meters on {times[0]}', level=1)

    output_text = (f'''
    Hi.
    I hope this message finds you well. I wanted to share the details of my recent {types_of_runs} for your analysis and feedback.

    **Run Details:**
    - Distance: {total_km_result} meters
    - Duration: {times[0]} to {times[-1]}

    **Average Pace per Kilometer:**
    - {res_Speed}

    **Average Heart Rate per Kilometer:**
    - {res_hr}

    **Average Cadence per Kilometer:**
    - {res_cd}

    **Average Altitude per Kilometer:**
    - {res_al}

    **Standard Deviation:**
    - Pace: {standard_deviation_pace} m/s
    - Heart Rate: {standard_deviation_heart_rate} bpm
    - Cadence: {standard_deviation_cadence} steps/min
    - Altitude: {standard_deviation_altitude} m

    I would greatly appreciate your insights and any recommendations you may have based on this data. Thank you in advance for your guidance.

    Best regards,

    Nima

    These are my answers after the run:
    Physical Comfort:
    How did your body feel during the run? Any specific areas of discomfort or pain?
    {Physical_Comfort_Response1}
    Did you experience any muscle tightness, soreness, or stiffness?
    {Physical_Comfort_Response2}


    Breathing and Cardiovascular Response:
    How was your breathing throughout the run? Was it comfortable or labored?
    {Breathing_and_Cardiovascular_Response1}
    Did you notice any changes in your heart rate, and if so, were they within a comfortable range?
    {Breathing_and_Cardiovascular_Response2}

    Energy Levels:
    How would you describe your energy levels during the run? Did you feel fatigued at any point?
    {Energy_Levels_Response}

    Mental State:
    What was your mental state like during the run? Were you focused, distracted, or perhaps in a flow state?
    {Mental_State_Response1}
    Did you experience any mental barriers or breakthroughs?
    {Mental_State_Response2}

    Running Form:
    Were you mindful of your running form? Did you notice any changes in your gait or posture?
    {Running_Form_Response1}
    Did you encounter any challenges related to your form?
    {Running_Form_Response2}

    Terrain and Environment:
    How did the terrain (e.g., flat, hilly, uneven) impact your running experience?
    {Terrain_and_Environment_Response1}
    Did the weather or environmental conditions affect your performance or enjoyment?
    {Terrain_and_Environment_Response2}

    Hydration and Nutrition:
    Did you feel adequately hydrated and fueled before and during the run?
    {Hydration_and_Nutrition_Response1}
    Did you experience any issues related to nutrition or hydration?
    {Hydration_and_Nutrition_Response2}

    Goal Achievement:
    Were you able to meet the goals you set for yourself during the run?
    {Goal_Achievement_Response1}
    How do you feel about your overall performance and progress?
    {Goal_Achievement_Response2}

    Recovery:
    How is your body feeling post-run? Any lingering discomfort or signs of quick recovery?
    {Recovery_Response1}
    Did you engage in any post-run stretching or recovery activities?
    {Recovery_Response2}

    Overall Satisfaction:
    On a scale of 1 to 10, how satisfied are you with your running experience today?
    {Overall_Satisfaction1}
    What aspects of the run brought you the most joy or fulfillment?
    {Overall_Satisfaction2}
    ''')

    # Add the output text to the document
    doc.add_paragraph(output_text)

    # Save the document
    doc_file_path = os.path.join(save_directory, f'{name} - {types_of_runs} - {times[0].split()[0]} - {total_km_result}.docx')
    doc.save(doc_file_path)

    print("Output has been printed to Word document:", doc_file_path)

if __name__ == "__main__":
    main()